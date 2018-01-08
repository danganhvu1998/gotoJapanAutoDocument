import zipfile
import shutil
import os
import sys
import json
import urllib.request, urllib.parse, urllib.error
from bs4 import BeautifulSoup
import ssl
import xml.etree.ElementTree as ET
import requests
import re
import openpyxl
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.compat import range
from openpyxl.utils import get_column_letter
from collections import Counter
import multiprocessing
import time
import codecs
import tempfile
import webbrowser
import docx
from docx import Document
from datetime import datetime
from tkinter import *
from tkinter import filedialog

strLogedIn = "no"
intlogedIn = 0
studentInfomation = dict()
student = dict()
USERNAME = str()
PASSWORD = str()
limitCol = 200
#destination = "/home/kyatodanhvu/Dropbox/VtmFileSystem/Python/"
destination = ""
desResult = str()
desFrom = str()
dataChecking = 0

################################################################################################################################################
###############################################                  AUTO RECOGNIZE                  ###############################################
def dataTaker(path):
	if("xlsx" in path):
		document = zipfile.ZipFile(path,'r')
		data = document.read('xl/sharedStrings.xml').decode()
	else:
		document = zipfile.ZipFile(path,'r')
		data = document.read('word/document.xml').decode()
	document.close()
	return(data)

def readingDocx(path):
	document = Document(path)
	ansText = str()
	for para in document.paragraphs:
		ansText = ansText+para.text
	tables = document.tables
	for table in tables:
		for row in table.rows:
			for cell in row.cells:
				for para in cell.paragraphs:
					ansText = ansText+para.text
	return(ansText)

def docxRecognize(fileIn, fileForm):
	dictResult = dict()
	#taking data
	dataInput = readingDocx(fileIn)
	dataForm = readingDocx(fileForm)
	#print(dataInput)
	#print(dataForm)

	#taking varName
	listVarName = list()
	workBook = load_workbook("HoSo_Chuan.xlsx")
	workSheet = workBook["khong_sua_ten"]
	for intCol in range(1,2000):
		valName = workSheet.cell(column=intCol, row=2).value
		if(valName == None): break
		listVarName.append(valName)

	#splitting info
	tempDataForm = dataForm
	for varName in listVarName:
		#print(varName, varName in tempDataForm)
		tempDataForm = tempDataForm.replace(varName,"#_#")
	#print(tempDataForm)
	#return
	listStr = tempDataForm.split("#_#")
	for varStr in listStr:
		#print(varStr)
		if(len(varStr)<1): continue
		dataForm = dataForm.replace(varStr,"#_#")
		dataInput = dataInput.replace(varStr,"#_#")
		#print("\""+varStr+"\"",len(dataForm))
	#print(dataForm)
	#print(dataInput)
	#return
	listVarName = dataForm.split("#_#")
	listVarData = dataInput.split("#_#")
	cnt = -1
	for varName in listVarName:
		cnt = cnt+1
		try:
			if(len(varName)>0 and dictResult.get(varName,"LLOOLLDAV")=="LLOOLLDAV"):
				varData = listVarData[cnt]
				if(varName!=varData):
					dictResult[varName]=varData
		except:
			break
	return(dictResult)

def xlsxRecognize(fileIn, fileForm):
	dictResult = dict()
	#taking data
	dataInput = dataTaker(fileIn)
	dataForm = dataTaker(fileForm)

	#taking varName
	listVarName = list()
	workBook = load_workbook("HoSo_Chuan.xlsx")
	workSheet = workBook["khong_sua_ten"]
	for intCol in range(1,2000):
		valName = workSheet.cell(column=intCol, row=2).value
		if(valName == None): break
		listVarName.append(valName)

	#splitting info
	tempDataForm = dataForm
	for varName in listVarName:
		#print(varName, varName in tempDataForm)
		tempDataForm = tempDataForm.replace(varName,"#_#")
	listStr = tempDataForm.split("#_#")
	for varStr in listStr:
		dataForm = dataForm.replace(varStr,"#_#")
		dataInput = dataInput.replace(varStr,"#_#")
	print(dataForm,dataInput)
	#return
	listVarName = dataForm.split("#_#")
	listVarData = dataInput.split("#_#")
	cnt = -1
	for varName in listVarName:
		cnt = cnt+1
		try:
			if(len(varName)>0 and dictResult.get(varName,"LLOOLLDAV")=="LLOOLLDAV"):
				varData = listVarData[cnt]
				dictResult[varName]=varData
		except:
			break
	return(dictResult)

def uploadStudent_AutoRead(upStudentID):
	workBook = load_workbook("HoSo_Chuan.xlsx")
	workSheet = workBook["khong_sua_ten"]
	###################################################
	#dirStudentIDTxt = "data_student/"+upStudentID+".txt"
	fhand = codecs.open("autoRecognizeInfo.txt", "r", "utf-8")
	data = "<"+upStudentID+">"

	for row in range(3,1000000):
		if(workSheet.cell(column=1, row=row).value==None): break
		if(str(workSheet.cell(column=1, row=row).value)!=upStudentID): continue
		for line in fhand:
			foundVar = 0
			#print(line)
			if("#EoF#" in line): break
			#print("What the fuck: ",line)
			studentVarName = line.split(":")[0].strip()
			studentVarData = line.split(":")[1].strip()
			for col in range (2,limitCol):
				if(workSheet.cell(column=col, row=2).value==None):
					foundCol = col
					break
				if(workSheet.cell(column=col, row=2).value==studentVarName):
					foundVar = 1
					if(str(workSheet.cell(column=col, row=row).value)!=studentVarData):
						if(str(workSheet.cell(column=col, row=row).value)=="None" and len(studentVarData)==0): continue
						data = data+studentVarName+"="+studentVarData+";"
						print("\tupdate",studentVarName," = ",studentVarData)
			if(foundVar==0):
				workSheet.cell(column=foundCol, row=2).value = studentVarName
				workSheet.cell(column=foundCol, row=row).value = studentVarData
				data = data+studentVarName+"="+studentVarData+";"
				print("\tupdate",studentVarName," = ",studentVarData)
		break

	if(data[-1]== ">"): 
		print("(SYS:)\tUpload finished! No change found!")
		fhand.close()
		return
	data = data+"--"+upStudentID+"--"
	#print(data)
	uploadServer(data)
	fhand.close()

def autoRecognize(auStudentID, dirForm, dirInput):

	#taking info	
	if("docx" in dirInput):
		#docx FIRSR
		dictData = docxRecognize(dirInput,dirForm)
	else:
		#xlsx SECOND
		dictData = xlsxRecognize(dirInput,dirForm)

	fhand = codecs.open("autoRecognizeInfo.txt", "w", "utf-8")
	for varName in dictData:
		fhand.write(varName+"\t:\t"+dictData[varName]+"\n")
	fhand.write("#EoF#\nYou always can create new data anyline above EoF line. (data : studentInfo)\nstudentInfo then will replace all data in student's document")
	fhand.close()
		


################################################################################################################################################

def updateZip(zipname, filename, data): 
    # generate a temp file
    tmpfd, tmpname = tempfile.mkstemp(dir=os.path.dirname(zipname))
    os.close(tmpfd)

    # create a temp copy of the archive without filename
    with zipfile.ZipFile(zipname, 'r') as zin:
        with zipfile.ZipFile(tmpname, 'w') as zout:
            zout.comment = zin.comment # preserve the comment
            for item in zin.infolist():
                if item.filename != filename:
                    zout.writestr(item, zin.read(item.filename))

    # replace with the temp archive
    os.remove(zipname)
    os.rename(tmpname, zipname)

    # now add filename with its new data
    with zipfile.ZipFile(zipname, mode='a', compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(filename, data)

def autoImportXlsx(path, intDirect):
	#print("==============",path)
	global desResult, desFrom, destination
	#copy file
	#urlFrom = desResult+path
	#urlTo = destination+path
	#shutil.copy2(urlFrom,urlTo)

	#import file
	document = zipfile.ZipFile(path,'r')
	data = document.read('xl/sharedStrings.xml').decode()

	fhand = codecs.open(desResult+"0fillInfo.txt","r","utf-8")
	for line in fhand:
		lineVarName = line.strip().split("~")[intDirect]
		lineVarData = line.strip().split("~")[1-intDirect]
		data = data.replace(lineVarName,lineVarData)
	document.close()

	updateZip(path,'xl/sharedStrings.xml',data)
	#document = zipfile.ZipFile(path,'a')
	#document.writestr('xl/sharedStrings.xml',data)
	#document.close()

	#cutf file
	#urlFrom = urlTo
	#urlTo = desResult+path
	#shutil.move(urlFrom,urlTo)

def autoImportDocx(path, intDirect):
	#print("==============",path)
	global desResult, desFrom, destination
	#copy file
	#urlFrom = desResult+path
	#urlTo = destination+path
	#shutil.copy2(urlFrom,urlTo)

	#import file
	document = zipfile.ZipFile(path,'r')
	data = document.read('word/document.xml').decode()

	fhand = codecs.open(desResult + "0fillInfo.txt", "r", "utf-8")
	for line in fhand:
		lineVarName = line.strip().split("~")[intDirect]
		lineVarData = line.strip().split("~")[1-intDirect]
		data = data.replace(lineVarName,lineVarData)
	document.close()

	updateZip(path, 'word/document.xml', data)
	#document = zipfile.ZipFile(path,'a')
	#document.writestr('document.xml',data)
	#document.close()
	#cutf file
	#urlFrom = urlTo
	#urlTo = desResult+path
	#shutil.move(urlFrom,urlTo)

def author():
	print("Vudz6298 :D")
	print("danganhvu1998@gmail.com")

def studentHtml(stStudentID):
	with requests.Session() as c:
		url="http://hr.vtmgroup.com.vn/user/login"
		#c.get(url)
		c.post(url, data = {'UserLogin[username]' : USERNAME, 'UserLogin[password]' : PASSWORD, 'UserLogin[rememberMe]' : '0'} )
		urlStudentID = "http://hr.vtmgroup.com.vn/students/"+stStudentID
		page = c.get(urlStudentID).content.decode()
		#print(page)
		if( "Xem chi tiết học viên" in page ) : return(page)
		return(0)

def defStudentInfo(deStudentID):
	#-----------Done Student Info-----------#
	#-----------Need to add StudentRelaJapan(Name, Rela, Address, Number)-----------#
	#-----------Info about relationship in Japan-----------#
	#-----------Add : 1 sponsor, 1 highschool, Mother+Father info(birth, job)-----------#
	countST = 0
	global studentInfomation
	global student
	page = studentHtml(deStudentID)
	if(page == 0): return(0)
	print("StudentFound")
	page = page.replace("Học sinh cấp 3 -&gt; Đại học","Học sinh cấp 3 lên Đại học")
	studentInfomation = re.findall('<tr class=".*"><th>.*</th><td>(.*)</td></tr>',page)
	for StudentInfo in studentInfomation:
		strStudentInfo = str(StudentInfo)
		if("000" in strStudentInfo) or ("span class" in strStudentInfo) or ("null" in strStudentInfo): studentInfomation[countST] = None
		countST = countST+1
	#print("wtf    ",studentInfomation)
	student['ID'] = studentInfomation[0]
	student['NaMe'] = studentInfomation[1]
	student['Gender'] = studentInfomation[2]
	student['Birthday'] = studentInfomation[3]
	student['CMNDNumber'] = studentInfomation[4]
	student['CMNDDate'] = studentInfomation[5]
	student['CMNDPlace'] = studentInfomation[6]
	student['Passport'] = studentInfomation[7]
	student['Number'] = studentInfomation[8]
	student['Email'] = studentInfomation[9]
	student['PSchoolName'] = studentInfomation[35]
	student['PSchoolData'] = studentInfomation[36]
	student['SSchoolName'] = studentInfomation[37]
	student['SSchoolDate'] = studentInfomation[38]
	student['HSchool1Name'] = studentInfomation[39]
	student['HSchool1Date'] = studentInfomation[40]
	student['University1Name'] = studentInfomation[41]
	student['University1Date'] = studentInfomation[42]
	student['University2Name'] = studentInfomation[43]
	student['University2Date'] = studentInfomation[44]
	student['Company1Name'] = studentInfomation[45]
	student['Company1Date'] = studentInfomation[46]
	student['Company2Name'] = studentInfomation[47]
	student['Company2Date'] = studentInfomation[48]
	student['Company3Name'] = studentInfomation[49]
	student['Company3Date'] = studentInfomation[50]
	student['Sponsor1Name'] = studentInfomation[51]
	student['Sponsor1Birthday'] = studentInfomation[52]
	student['Sponsor1Number'] = studentInfomation[53]
	student['Sponsor1JobName'] = studentInfomation[54]
	student['Sponsor1DNN'] = studentInfomation[55]
	student['Sponsor1JobPlace'] = studentInfomation[56]
	student['Rela'] = studentInfomation[57]
	student['NguyenQuan'] = studentInfomation[59]
	student['ThuongChu'] = studentInfomation[60]
	student['BirthPlace'] = studentInfomation[61]
	student['Address'] = studentInfomation[62]
	student['FatherName'] = studentInfomation[63]
	student['MotherName'] = studentInfomation[64]

def createStudentTxt(crStudentID):#open HoSo_Chuan
	global dataChecking
	dirStudentIDTxt = "data_student/"+crStudentID+".txt"
	fhand = codecs.open(dirStudentIDTxt, "w", "utf-8")
	##################################################
	##############Frist, read HoSo_Chuan##############
	##################################################
	workBook = load_workbook("HoSo_Chuan.xlsx")
	workSheet = workBook["khong_sua_ten"]	
	##################################################
	for row in range(3,1000000):
		if(str(workSheet.cell(column=1, row=row).value)!=crStudentID) and (workSheet.cell(column=1, row=row).value!=None): continue
		#print("ROW = ",row)
		workSheet.cell(column=1, row=row).value=crStudentID #none or the right row are the same
		# now workSheet.cell(column=1, row=row).value == crStudentID
		for col in range(2,limitCol):# do not print ID
			if(workSheet.cell(column=col, row=2).value==None):break #read all col
			dataWrite=str(workSheet.cell(column=col, row=2).value)+"\t:\t"+str(workSheet.cell(column=col, row=row).value)+"\n"
			fhand.write(dataWrite)
		workBook.save(filename = "HoSo_Chuan.xlsx")
		break

	fhand.write("#EoF#")
	fhand.close()
	# now time for check it
	#if(defStudentInfo(crStudentID)==0):
	#	fhand = codecs.open(dirStudentIDTxt, "a", "utf-8")
	#	fhand.write("#EoF#\nData from HR\nNo Data form HR\n")
	#	fhand.close()
	#	return
	dataFull = ""
	dataLack = ""
	dataDiff = ""
	fhand = codecs.open(dirStudentIDTxt, "r", "utf-8")
	for line in fhand:
		if("#EoF#" in line): break
		studentVarName = line.split(":")[0].strip()
		studentVarData = line.split(":")[1].strip()
		try:
			studentVarName = re.findall("student(.*)$",studentVarName)[0]
			studentVarNameData = student.get(studentVarName,"None").strip()
			#print(studentVarName+"+"+studentVarData+"+"+studentVarNameData+"+")
			if(studentVarData=="None" or len(studentVarData)==0):
				if(studentVarNameData=="None" or len(studentVarNameData)==0):
					dataLack = dataLack+"student"+studentVarName+"\t:\t\n"
				else:
					dataFull = dataFull+"student"+studentVarName+"\t:\t"+studentVarNameData+"\n"
			else:
				if(studentVarNameData=="None" or len(studentVarNameData)==0):
					dataFull = dataFull+"student"+studentVarName+"\t:\t"+studentVarData+"\n"
				else:
					dataFull = dataFull+"student"+studentVarName+"\t:\t"+studentVarData+"\n"
					if(studentVarNameData!=studentVarData):
						dataDiff = dataDiff+"student"+studentVarName+"\t:\t\""+studentVarData+"\" and \""+studentVarNameData+"\"\n"
		except:
			studentVarName = line.split(":")[0].strip()
			if(studentVarData=="None" or len(studentVarData)==0):
				dataLack = dataLack+studentVarName+"\t:\t\n"
			else:
				dataFull = dataFull+studentVarName+"\t:\t"+studentVarData+"\n"
	fhand.close()
	fhand = codecs.open(dirStudentIDTxt, "w", "utf-8")
	fhand.write(dataFull)
	fhand.write(dataLack)
	fhand.write("#EoF#\nYou always can create new data anyline above EoF line. (data : studentInfo)\nstudentIndo then will replace all data in student's document\nDifferent data between SERVER and HR(Data Server always has higher priority):\n")
	fhand.write(dataDiff)
	if(dataChecking==1):
		ErrLog = codecs.open("ErrLog.txt", "a", "utf-8")
		ErrLog.write(dataDiff)
		ErrLog.write("#####\n ### \n  #  \n ### \n#####\n\n")
		ErrLog.close()
		print(dataDiff)
		print("#####\n ### \n  #  \n ### \n#####\n\n")
	fhand.close()

def updateData():
	url = "http://devhr.vtmgroup.com.vn/api_test.php?"
	keyMd5 = "name=takemd51432j41j23k515b2b345vgfjhg8daf70ds89fuajljh24332jkq3eafkdufyaosfhajwenqkha0f807ds0af87a0dsf87asd0f8"
	keyDown = "name=download1432k4g12gj5f1g5kh1b13hg4hk3g1k4h31b4132v4c412gfd41d3sg4d5ddf8a879af8ds7fa98708b70v7z0vc87z70f&dulieu="
	urlMd5 = url + keyMd5
	urlDown =  url + keyDown
	with requests.Session() as c:
		page = c.get(urlMd5).content.decode()
	fhand = open("webServer/md5.txt")
	inp = fhand.read()
	fhand.close()
	pageServerHad = int(re.findall("=(.*)=",page)[0])
	intStart = max(0,pageServerHad-500)

	fhand = open("webServer/md5.txt","a")
	#fhandData = open("webServer/dataUpdate.txt","a")
	for pageCount in range(intStart, pageServerHad):
		strFind = "\("+str(pageCount)+"\).*"+"\("+str(pageCount)+"\)"
		strMd5 = re.findall(strFind,page)[0]
		if(strMd5 not in inp):
			fhand.write(strMd5)
			#download the updade
			urlDownPage =  urlDown+str(pageCount)
			with requests.Session() as c:
				pageData = c.get(urlDownPage).content.decode()
			addData = re.findall("<data>(.*)<data>",pageData)[0]
			#fhandData.write(addData)
			updateExel(addData,pageCount)


	urlDownPage =  urlDown+str(pageServerHad)
	with requests.Session() as c:
		pageData = c.get(urlDownPage).content.decode()
	addData = re.findall("<data>(.*)<data>",pageData)[0]
	#fhandData.write(addData)
	updateExel(addData,-1)#open HoSo_Chuan

def updateExel(updateData,intBackUp):#open HoSo_Chuan
	print("running")
	print(updateData)
	print("running")
	if(len(updateData)==0): return
	idList = re.findall("<([0-9]+)>",updateData)
	idList.sort()
	idLIST = list()
	rowDataList = list()
	cellDataList = list()
	idLIST.append(idList[0])
	for ID in idList:
		if(ID!=idLIST[-1]): idLIST.append(ID)
	#print(idLIST)
	##############################################
	workBook = load_workbook("HoSo_Chuan.xlsx")
	workSheet = workBook["khong_sua_ten"]
	##############################################
	updateRow = 0
	for ID in idLIST:
		for row in range (3,1000000):
			if workSheet.cell(column=1, row=row).value==None:
				updateRow=row
				workSheet.cell(column=1, row=row).value=ID
				break
			elif (str(workSheet.cell(column=1, row=row).value)==ID):
				updateRow=row
				break
		rowData = "<"+ID+">"+"(.*?)"+"--"+ID+"--"
		rowDataList = re.findall(rowData, updateData)
		#print("RowData",rowData,len(rowDataList))
		#print("\n\n\n\n\n\n\n\n\n")
		#for rowData in rowDataList:
		#	print(ID, rowData)
		for rowData in rowDataList:
			cellDataList = rowData.split(";")
			for cellData in cellDataList:
				try:
					#print("cellData =",cellData)
					cellVarName = cellData.split("=")[0]
					cellVarData = cellData.split("=")[1]
					#print(ID,updateRow,cellVarName,cellVarData)
					#print(cellVarName," : ",cellVarData)
					for col in range(2,limitCol):
						if(workSheet.cell(column=col, row=2).value==None):
							updateCol = col
							workSheet.cell(column=col, row=2).value = cellVarName
							break
						elif(workSheet.cell(column=col, row=2).value==cellVarName):
							updateCol = col
							break
					workSheet.cell(column=updateCol, row=updateRow).value = cellVarData
				except:
					#print("\nrun fucking ning\n")
					if(intBackUp!=-1):
						strNameBackUp = "result/backUp/HoSo_BackUp"+str(intBackUp)+".xlsx"
						workBook.save(strNameBackUp)
	workBook.save(filename = "HoSo_Chuan.xlsx")
	print("Updated!",datetime.now())

def uploadStudent(upStudentID):#open HoSo_Chuan
	workBook = load_workbook("HoSo_Chuan.xlsx")
	workSheet = workBook["khong_sua_ten"]
	###################################################
	dirStudentIDTxt = "data_student/"+upStudentID+".txt"
	fhand = codecs.open(dirStudentIDTxt, "r", "utf-8")
	data = "<"+upStudentID+">"

	for row in range(3,1000000):
		if(workSheet.cell(column=1, row=row).value==None): break
		if(str(workSheet.cell(column=1, row=row).value)!=upStudentID): continue
		for line in fhand:
			foundVar = 0
			if("#EoF#" in line): break
			#print(line)
			studentVarName = line.split(":")[0].strip()
			studentVarData = line.split(":")[1].strip()
			for col in range (2,limitCol):
				if(workSheet.cell(column=col, row=2).value==None):
					foundCol = col
					break
				if(workSheet.cell(column=col, row=2).value==studentVarName):
					foundVar = 1
					if(str(workSheet.cell(column=col, row=row).value)!=studentVarData):
						if(str(workSheet.cell(column=col, row=row).value)=="None" and len(studentVarData)==0): continue
						data = data+studentVarName+"="+studentVarData+";"
						print("\tupdate",studentVarName," = ",studentVarData)
			if(foundVar==0):
				workSheet.cell(column=foundCol, row=2).value = studentVarName
				workSheet.cell(column=foundCol, row=row).value = studentVarData
				data = data+studentVarName+"="+studentVarData+";"
				print("\tupdate",studentVarName," = ",studentVarData)
		break

	if(data[-1]== ">"): 
		print("(SYS:)\tUpload finished! No change found!")
		fhand.close()
		return
	data = data+"--"+upStudentID+"--"
	uploadServer(data)
	fhand.close()

def uploadServer(data):#rewrite uploadRow->uploadFile
	#print(data)
	url = "http://devhr.vtmgroup.com.vn/api_test.php?"
	key = "name=uploadDadAfasdNGtrwtAetNyreHy5V767U54TiyuHjfEhzvcjfHdsAbNcnbdhmvDfshSOf4g757757aMfaEzvB!74567674$774%"
	data = "&dulieu="+data+"---"
	url = url+key+data
	if(len(data)>5000):
		print("(SYS:)\tToo Much data. Try to send it in two line. By the ways, you so should not send this much data to server!")
		return
	with requests.Session() as c:
		page = c.get(url).content.decode()
	print("(SYS:)\t",re.findall("\(result\)(.*)\(result\)",page)[0])

def exportFile(exStudentID, exGroup):#rewrite - last one
	print(exStudentID,exGroup)
	#sreturn
	##############################################
	global desResult, desFrom, destination
	workBook = load_workbook("HoSo_Chuan.xlsx")
	workSheet = workBook["khong_sua_ten"]

	for row in range(3,1000000):
		if (str(workSheet.cell(column=1, row=row).value)==exStudentID):
			exStudentName = workSheet.cell(column=2, row=row).value
			exRow = row
			break
		if (str(workSheet.cell(column=1, row=row).value)==None):
			print("(SYS:)\tCan not find student with ID = ", exStudentID)
			return
	##############################################
	exStudentName = exStudentName.replace(" ","-")
	desFrom = exGroup+"/"
	exGroup = exGroup.split("/")[-1]
	print(exGroup)
	desResult = destination+"result/"+exStudentID+"_"+exStudentName+"_"+exGroup+"/"
	if not os.path.exists(desResult):
		print("(SYS:)\tFound no file. Auto creating ...")
		fileDesFrom = os.listdir( desFrom )
		os.makedirs(desResult)
		fhand = codecs.open(desResult + "0fillInfo.txt", "w", "utf-8")
		for col in range(2,limitCol):
			if(workSheet.cell(column=col, row=2).value==None): break
			if(workSheet.cell(column=col, row=exRow).value==None): continue
			fhand.write(str(workSheet.cell(column=col, row=2).value)+"~"+str(workSheet.cell(column=col, row=exRow).value)+"\n")
		fhand.close()
		for file in fileDesFrom:
			urlFrom = desFrom+file
			urlTo = desResult+file
			shutil.copy2(urlFrom,urlTo)
		for file in fileDesFrom:
			if("~" in file): continue
			if(".docx" in file):
				autoImportDocx(desResult+file, 0)
			elif(".xlsx" in file):
				autoImportXlsx(desResult+file, 0)
	else:
		print("(SYS:)\tWorking ... Remember, if you want to add some more document, add it directly to", desResult)
		fileDesResult = os.listdir( desResult )
		for file in fileDesResult:
			if("~" in file): continue
			if(".docx" in file):
				autoImportDocx(desResult+file, 1)
			elif(".xlsx" in file):
				autoImportXlsx(desResult+file, 1)

		fhand = codecs.open(desResult+"0fillInfo.txt","w","utf-8")
		for col in range(2,limitCol):
			if(workSheet.cell(column=col, row=2).value==None): break
			if(workSheet.cell(column=col, row=exRow).value==None): continue
			#print(str(workSheet.cell(column=col, row=2).value)+"~"+str(workSheet.cell(column=col, row=exRow).value)+"\n")
			fhand.write(str(workSheet.cell(column=col, row=2).value)+"~"+str(workSheet.cell(column=col, row=exRow).value)+"\n")
		fhand.close()

		for file in fileDesResult:
			if("~" in file): continue
			if(".docx" in file):
				autoImportDocx(desResult+file, 0)
			elif(".xlsx" in file):
				autoImportXlsx(desResult+file, 0)

def userRewriteStudent(usStudentID):
	try:
		os.startfile("data_student\\"+usStudentID+".txt")		
	except:
		webbrowser.open("data_student/"+usStudentID+".txt")

def seeStudent(seStudentID):
	updateData()
	createStudentTxt(seStudentID)
	while(1):
		updateData()
		action = input("(Student INFO) What do you need? (Name = "+student["NaMe"]+" ) \n\t+reco : Auto read info from document\n\t+see : See and edit student info\n\t+upload : upload student info to server\n\t+export : Export student file\n\t+exit : quit\n")
		action = action.lower().strip()
		if(action == "see"):
			createStudentTxt(seStudentID)
			userRewriteStudent(seStudentID)
			action = input("do you wish to upload it to server(Save file ID.txt first)?(y/n)").lower().strip()
			if(action == "y"):
				uploadStudent(seStudentID)
		elif(action == "upload"):
			action = input("You sure that already saved ID.txt file?(y/n)").lower().strip()
			if(action == "y"):
				uploadStudent(seStudentID)
		elif(action == "reco"):
			autoRecognize(seStudentID)
		elif(action == "export"):
			print("Please save all the data in ID.txt")
			uploadStudent(seStudentID)
			updateData()
			print("Export with which school document? Type the number or the name exactly.")
			cnt = 1
			ckk = 0
			fileList = os.listdir("data/")
			for file in fileList:
				print("\t+",cnt,file)
				cnt=cnt+1
			studentSchoolDocument = input().lower().strip()
			try:
				intFile = int(studentSchoolDocument)-1
				exportFile(seStudentID,fileList[intFile])
				print("(SYS:)\tDone! You can find all files in result")
			except:
				for file in fileList:
					if(file==studentSchoolDocument):
						ckk = 1
						exportFile(seStudentID,file)
						print("(SYS:)\tDone! You can find all files in result")
						break
				if(ckk==0): print("(SYS:)\tErr: No school found!\n")
		elif(action == "exit"):
			return
		else:
			print("(SYS:)\tBad command! I cannot understand!")
	
def reviewALL():
	global dataChecking
	MASTERPASSWORD = input("ENTER MASTERPASSWORD:")
	#if(MASTERPASSWORD != "123"): return
	if(MASTERPASSWORD != "GOTOJAPAN_GTJP_ahihidongok_!#@_123"): return
	if( input("Do you understand that this performant costs a lot of time and does require good Internet Connection?(y/n)")=="y"):
		print("(SYS:)\trunning")
		ErrLog = codecs.open("ErrLog.txt", "w", "utf-8")
		ErrLog.close()
		dataChecking = 1
		workBook = load_workbook("HoSo_Chuan.xlsx")
		workSheet = workBook["khong_sua_ten"]
		for row in range(3,1000000):
			ErrLog = codecs.open("ErrLog.txt", "a", "utf-8")
			reStudentID = str(workSheet.cell(column=1, row=row).value)
			if(reStudentID=="None"): break
			reStudentName = str(workSheet.cell(column=2, row=row).value)
			print("(SYS:)\tchecking studentName =", reStudentName, "with ID =", reStudentID)
			ErrLog = codecs.open("ErrLog.txt", "a", "utf-8")
			ErrLog.write("studentName = "+reStudentName+" with ID = "+reStudentID+"\n")
			ErrLog.close()
			createStudentTxt(reStudentID)
		dataChecking = 0
		ErrLog.close()

class AppGTJPLogin(Frame):
	def __init__(self, master):
		Frame.__init__(self, master)
		self.master = master
		self.master.title("Login HR")
		#self.master.geometry("400x300")
		Label(master, text="Username: ").grid(row=0)
		Label(master, text="Password: ").grid(row=1)
		self.e1 = Entry(master)
		self.e2 = Entry(master)
		self.e1.grid(row=0, column=1)
		self.e2.grid(row=1, column=1)
		Button(master, text='Quit', command=self.master.quit).grid(row=3, column=0, sticky=W, pady=4)
		Button(master, text='Login', command=self.userLogin).grid(row=3, column=1, sticky=W, pady=4)

	def userLogin(self):
		global intLogedIn, strLogedIn, USERNAME, PASSWORD
		USERNAME = self.e1.get()
		PASSWORD = self.e2.get()
		with requests.Session() as c:
			url="http://hr.vtmgroup.com.vn/user/login"
			page = (c.post(url, data = {'UserLogin[username]' : USERNAME, 'UserLogin[password]' : PASSWORD, 'UserLogin[rememberMe]' : '0'} )).content.decode()
			if "Hồ sơ cá nhân"  not in page:
				print("Kiểm tra lại Usernam hoặc Password")
				self.e1.delete(0,END)
				self.e2.delete(0,END)
			else :
				print("ĐĂNG NHẬP THÀNH CÔNG")
				intLogedIn = 1
				strLogedIn ="yes"
				self.e1.delete(0,END)
				self.e2.delete(0,END)
				self.master.quit()

class AppGTJPStudent(Frame):
	def __init__(self, master):
		Frame.__init__(self, master)
		####################################ForderReader####################################
		#self.folderData = Toplevel()
		#self.folderData.title("Choose school or group of files ...")
		####################################ForderReader####################################
		self.master = master
		self.master.title("Auto Document")
		self.strStudentID = "start"
		Button(master, text='Quit', command=self.master.quit).pack(side=LEFT)
		Label(master, text="studentID: ").pack(side=LEFT)
		self.e1 =  Entry(master)
		self.e1.pack(side=LEFT)
		#Button(master, text='Show', command=self.showStudentInfo).pack(side=BOTTOM)
		self.studentText = Text(master, height=40, width=75)
		self.studentText.pack(side=BOTTOM)
		Button(master, text='Show', command=self.showStudentInfo).pack(side=LEFT)
		Button(master, text='Save and Update', command=self.saveAndUpdate).pack(side=LEFT)
		Button(master, text='Export',command=self.studentExport).pack(side=LEFT)
		Button(master, text='Auto Read',command=self.autoRead).pack(side=LEFT)

	def showStudentInfo(self):
		updateData()
		self.studentText.delete('1.0', END)
		self.strStudentID = str(self.e1.get())
		if( defStudentInfo(self.strStudentID)==0 ):
			self.studentText.insert(END,"Student Not Found! Recheck the ID or your Internet Connection")
			return
		createStudentTxt(self.strStudentID)
		dirStudentIDTxt = "data_student/"+self.strStudentID+".txt"
		fhand = codecs.open(dirStudentIDTxt, "r", "utf-8")
		for line in fhand:
			#print(line)
			if( len(line.strip())==0 ): continue
			self.studentText.insert(END,line)
		fhand.close()

	def studentExport(self):
		folderData = filedialog.askdirectory(initialdir = "data/",title = "Select school")
		exportFile(str(self.e1.get()), folderData)

	def saveAndUpdate(self):
		if( self.strStudentID!=str(self.e1.get()) ):
			self.studentText.delete('1.0', END)
			self.showStudentInfo()
			return
		studentData = self.studentText.get("1.0",'end-1c')
		if("Student Not Found! Recheck the ID or your Internet Connection" in studentData):
			return
		dirStudentIDTxt = "data_student/"+self.strStudentID+".txt"
		fhand = codecs.open(dirStudentIDTxt, "w", "utf-8")
		fhand.write(studentData)
		fhand.close()
		uploadStudent(self.strStudentID)
		updateData()

	def autoRead(self):
		if( self.strStudentID!=str(self.e1.get()) ):
			self.studentText.delete('1.0', END)
			self.showStudentInfo()
		studentData = self.studentText.get("1.0",'end-1c')
		if("Student Not Found! Recheck the ID or your Internet Connection" in studentData):
			return
		studentDocument =  filedialog.askopenfilename( initialdir = "autoRecognize/", title = "Select student's document ...",filetypes = (("Word","*.docx"),("Exel","*.xlsx")) )
		studentForm =  filedialog.askopenfilename( initialdir = "data/", title = "Select form ...",filetypes = (("Word","*.docx"),("Exel","*.xlsx")) )
		if(".docx" in studentDocument) and (".docx" not in studentForm):
			self.studentText.delete('1.0', END)
			self.studentText.insert(END, "Err: DOCX with XLSX")
			return
		if(".xlsx" in studentDocument) and (".xlsx" not in studentForm):
			self.studentText.delete('1.0', END)
			self.studentText.insert(END, "Err: DOCX with XLSX")
		print(studentDocument+"\n"+studentForm)
		autoRecognize(self.strStudentID,studentForm,studentDocument)
		####recognized, update info to self.studentText
		self.studentText.delete('1.0', END)
		fhand = codecs.open("autoRecognizeInfo.txt", "r", "utf-8")
		for line in fhand:
			#print(line)
			if( len(line.strip())==0 ): continue
			self.studentText.insert(END,line)
		fhand.close()

root = Tk()
app = AppGTJPLogin(root)
root.mainloop()
print(intLogedIn,strLogedIn)
if(intLogedIn==1 and strLogedIn=="yes"):
	Root =Tk()
	app = AppGTJPStudent(Root)
	Root.mainloop()
